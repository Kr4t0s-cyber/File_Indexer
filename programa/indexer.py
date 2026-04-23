#!/usr/bin/env python3
"""
File Indexer - Cross-platform file indexing and search application
Supports: PDF, DOCX, ODT, TXT, RTF, CSV, HTML, XML, MD and ZIP archives
Runs on macOS, Windows and Linux.
"""

import os
import sys
import sqlite3
import zipfile
import threading
import time
import json
import hashlib
import traceback
import platform
import subprocess
from pathlib import Path
from datetime import datetime
from io import BytesIO, StringIO
import re
import chardet

# Flask
from flask import Flask, request, jsonify, render_template_string, send_from_directory

# Text extractors
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfpage import PDFPage
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from odf.opendocument import load as odf_load
    from odf.text import P
    from odf import teletype
    HAS_ODF = True
except ImportError:
    HAS_ODF = False

# Word COM automation (only on Windows com Word instalado)
try:
    import pythoncom  # noqa: F401
    import win32com.client  # noqa: F401
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False


app = Flask(__name__)

DB_PATH = os.path.expanduser("~/.file_indexer/index.db")
os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)

# Global state
indexing_state = {
    "running": False,
    "progress": 0,
    "total": 0,
    "current_file": "",
    "errors": [],
    "done": False,
    "start_time": None,
    "indexed_count": 0,
    "skipped_count": 0,
}

SUPPORTED_EXTENSIONS = {
    '.pdf', '.docx', '.doc', '.odt', '.ods', '.odp',
    '.txt', '.text', '.rtf', '.csv', '.tsv',
    '.html', '.htm', '.xml', '.md', '.markdown',
    '.json', '.log', '.py', '.js', '.ts', '.css',
    '.java', '.c', '.cpp', '.h', '.rs', '.go', '.rb', '.php',
    '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf',
    '.tex', '.rst', '.nfo', '.srt', '.vtt',
}

ARCHIVE_EXTENSIONS = {'.zip', '.cbz'}


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn


def init_db():
    with get_db() as conn:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                path TEXT NOT NULL,
                filename TEXT NOT NULL,
                extension TEXT,
                size INTEGER,
                modified_at TEXT,
                indexed_at TEXT,
                hash TEXT,
                inside_archive TEXT,
                archive_path TEXT,
                word_count INTEGER DEFAULT 0,
                error TEXT
            );
            CREATE UNIQUE INDEX IF NOT EXISTS idx_files_path ON files(path);
            
            CREATE VIRTUAL TABLE IF NOT EXISTS file_content USING fts5(
                path UNINDEXED,
                content,
                tokenize='unicode61 remove_diacritics 1'
            );
            
            CREATE TABLE IF NOT EXISTS indexed_dirs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                path TEXT UNIQUE NOT NULL,
                indexed_at TEXT,
                file_count INTEGER DEFAULT 0
            );
        """)


class PasswordProtectedError(Exception):
    """Raised when a file is encrypted / password-protected and cannot be read."""
    pass


def _is_ooxml_encrypted(file_bytes):
    """A .docx / .xlsx / .pptx file is a ZIP (starts with 'PK'). When password
    protected, it is instead wrapped in an OLE compound file. We detect that
    by checking the magic bytes at the start of the file."""
    if not file_bytes or len(file_bytes) < 8:
        return False
    return file_bytes[:8] == b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'


def _is_doc_encrypted_heuristic(file_bytes):
    """Heuristic for .doc encryption: password-protected Office files contain
    specific stream names in their OLE structure ('EncryptedPackage',
    'EncryptionInfo'). Looking for the UTF-16LE-encoded names catches the
    common cases without needing to parse OLE properly."""
    if not file_bytes or len(file_bytes) < 512:
        return False
    head = file_bytes[:65536]
    markers = (
        b'E\x00n\x00c\x00r\x00y\x00p\x00t\x00e\x00d\x00P\x00a\x00c\x00k\x00a\x00g\x00e',
        b'E\x00n\x00c\x00r\x00y\x00p\x00t\x00i\x00o\x00n\x00I\x00n\x00f\x00o',
    )
    return any(m in head for m in markers)


def _extract_doc_via_word(filepath):
    """Extract text from a .doc file using Microsoft Word via COM automation.
    Requires Windows + Microsoft Word + pywin32.
    Returns extracted text, or "" if Word couldn't read it.
    Raises PasswordProtectedError if the file needs a password."""
    if not HAS_WIN32COM or not filepath:
        return ""
    import pythoncom
    import win32com.client
    word = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        # DisplayAlerts = 0 (wdAlertsNone) — suppress dialogs
        word.DisplayAlerts = 0
        try:
            # Passing a dummy password avoids the interactive password prompt.
            # If the doc actually needs a password, Word raises an error we
            # can catch below.
            doc = word.Documents.Open(
                os.path.abspath(filepath),
                ReadOnly=True,
                AddToRecentFiles=False,
                ConfirmConversions=False,
                Visible=False,
                PasswordDocument="__NO_PASSWORD__",
            )
        except Exception as e:
            msg = str(e).lower()
            if 'password' in msg or 'senha' in msg or 'protected' in msg:
                raise PasswordProtectedError()
            return ""
        try:
            text = doc.Content.Text or ""
        finally:
            doc.Close(SaveChanges=False)
        return text
    except PasswordProtectedError:
        raise
    except Exception:
        return ""
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _extract_doc_fallback(file_bytes):
    """Best-effort text extraction from a .doc binary without Word.
    Pulls runs of printable ASCII and UTF-16LE strings. Not perfect but
    enough for full-text search (keywords)."""
    if not file_bytes:
        return ""
    pieces = []
    # ASCII runs (length >= 4)
    pieces.extend(
        m.decode('latin-1', errors='replace')
        for m in re.findall(rb'[\x20-\x7E\xC0-\xFF]{4,}', file_bytes)
    )
    # UTF-16LE runs: char followed by null byte, repeated (length >= 4 chars)
    utf16_chunks = re.findall(rb'(?:[\x20-\x7E\xC0-\xFF]\x00){4,}', file_bytes)
    for chunk in utf16_chunks:
        try:
            pieces.append(chunk.decode('utf-16-le', errors='replace'))
        except Exception:
            pass
    text = '\n'.join(pieces)
    # Collapse crazy whitespace runs that binary parsing produces
    text = re.sub(r'[ \t]{3,}', '  ', text)
    return text


def extract_text_from_file(filepath, extension, file_bytes=None):
    """Extract text content from various file types."""
    try:
        if file_bytes is None:
            with open(filepath, 'rb') as f:
                file_bytes = f.read()

        ext = extension.lower()

        # PDF
        if ext == '.pdf' and HAS_PDF:
            try:
                text = pdf_extract_text(BytesIO(file_bytes))
                return text or ""
            except Exception:
                return ""

        # DOCX (formato moderno — ZIP com XML)
        if ext == '.docx' and HAS_DOCX:
            if _is_ooxml_encrypted(file_bytes):
                raise PasswordProtectedError()
            try:
                doc = DocxDocument(BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs]
                # Also extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                return '\n'.join(paragraphs)
            except PasswordProtectedError:
                raise
            except Exception:
                pass

        # DOC (formato binario antigo — Word 97-2003)
        if ext == '.doc':
            # Deteccao rapida de senha: arquivos protegidos tem marcadores
            # especificos na estrutura OLE
            if _is_doc_encrypted_heuristic(file_bytes):
                raise PasswordProtectedError()
            # Preferencia: usar Word via COM (se disponivel, resultado e perfeito)
            if HAS_WIN32COM and filepath:
                text = _extract_doc_via_word(filepath)  # pode lancar PasswordProtectedError
                if text:
                    return text
            # Fallback pure-Python: extracao de strings do binario
            return _extract_doc_fallback(file_bytes)

        # ODT / ODS / ODP
        if ext in ('.odt', '.ods', '.odp') and HAS_ODF:
            try:
                doc = odf_load(BytesIO(file_bytes))
                texts = teletype.extractText(doc.body)
                return texts
            except Exception:
                pass

        # RTF - basic stripping
        if ext == '.rtf':
            try:
                text = file_bytes.decode('latin-1', errors='replace')
                # Strip RTF control codes
                text = re.sub(r'\\[a-z]+\d* ?', ' ', text)
                text = re.sub(r'[{}\\]', ' ', text)
                return text
            except Exception:
                pass

        # Plain text formats
        text_exts = {
            '.txt', '.text', '.csv', '.tsv', '.md', '.markdown',
            '.html', '.htm', '.xml', '.json', '.log', '.py', '.js',
            '.ts', '.css', '.java', '.c', '.cpp', '.h', '.rs', '.go',
            '.rb', '.php', '.yaml', '.yml', '.toml', '.ini', '.cfg',
            '.conf', '.tex', '.rst', '.nfo', '.srt', '.vtt',
        }
        if ext in text_exts:
            # Detect encoding
            detected = chardet.detect(file_bytes[:10000])
            encoding = detected.get('encoding') or 'utf-8'
            try:
                return file_bytes.decode(encoding, errors='replace')
            except Exception:
                return file_bytes.decode('utf-8', errors='replace')

    except PasswordProtectedError:
        # Re-raise to let the caller skip this file with a proper message
        raise
    except Exception as e:
        return ""

    return ""


def compute_hash(file_bytes):
    return hashlib.md5(file_bytes[:65536]).hexdigest()


def index_file_entry(conn, filepath, filename, extension, size, modified_at, content, inside_archive=None, archive_path=None):
    """Insert or update a file in the index."""
    word_count = len(content.split()) if content else 0
    now = datetime.now().isoformat()

    # Unique path key
    path_key = f"{archive_path}::{filepath}" if archive_path else filepath

    conn.execute("""
        INSERT OR REPLACE INTO files 
        (path, filename, extension, size, modified_at, indexed_at, inside_archive, archive_path, word_count)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (path_key, filename, extension, size, modified_at, now, inside_archive, archive_path, word_count))

    # FTS - delete old, insert new
    conn.execute("DELETE FROM file_content WHERE path = ?", (path_key,))
    if content and content.strip():
        conn.execute("INSERT INTO file_content (path, content) VALUES (?, ?)", (path_key, content))


def collect_files(directory):
    """Collect all indexable files including inside zips."""
    files = []
    for root, dirs, filenames in os.walk(directory):
        # Skip hidden directories
        dirs[:] = [d for d in dirs if not d.startswith('.')]
        for fname in filenames:
            if fname.startswith('.'):
                continue
            fpath = os.path.join(root, fname)
            ext = Path(fname).suffix.lower()
            if ext in SUPPORTED_EXTENSIONS:
                files.append(('file', fpath, fname, ext))
            elif ext in ARCHIVE_EXTENSIONS:
                files.append(('archive', fpath, fname, ext))
    return files


def do_indexing(directory):
    global indexing_state
    indexing_state["running"] = True
    indexing_state["done"] = False
    indexing_state["errors"] = []
    indexing_state["indexed_count"] = 0
    indexing_state["skipped_count"] = 0
    indexing_state["start_time"] = time.time()
    indexing_state["current_file"] = "Coletando arquivos..."

    try:
        items = collect_files(directory)
        indexing_state["total"] = len(items)
        indexing_state["progress"] = 0

        conn = get_db()

        for i, item in enumerate(items):
            if not indexing_state["running"]:
                break

            kind = item[0]
            filepath = item[1]
            fname = item[2]
            ext = item[3]

            indexing_state["progress"] = i + 1
            indexing_state["current_file"] = fname

            try:
                stat = os.stat(filepath)
                size = stat.st_size
                modified_at = datetime.fromtimestamp(stat.st_mtime).isoformat()

                if kind == 'file':
                    with open(filepath, 'rb') as f:
                        file_bytes = f.read()
                    try:
                        content = extract_text_from_file(filepath, ext, file_bytes)
                    except PasswordProtectedError:
                        indexing_state["errors"].append(f"{fname}: protegido por senha — ignorado")
                        indexing_state["skipped_count"] += 1
                        # Commit every 50 files
                        if i % 50 == 0:
                            conn.commit()
                        continue
                    index_file_entry(conn, filepath, fname, ext, size, modified_at, content)
                    indexing_state["indexed_count"] += 1

                elif kind == 'archive':
                    try:
                        with zipfile.ZipFile(filepath, 'r') as zf:
                            for zinfo in zf.infolist():
                                if zinfo.is_dir():
                                    continue
                                zname = zinfo.filename
                                zext = Path(zname).suffix.lower()
                                if zext not in SUPPORTED_EXTENSIONS:
                                    continue
                                try:
                                    zdata = zf.read(zinfo.filename)
                                    try:
                                        zcontent = extract_text_from_file(None, zext, zdata)
                                    except PasswordProtectedError:
                                        indexing_state["errors"].append(f"{zname} (in {fname}): protegido por senha — ignorado")
                                        indexing_state["skipped_count"] += 1
                                        continue
                                    zfname = Path(zname).name
                                    zsize = zinfo.file_size
                                    zmod = datetime(*zinfo.date_time).isoformat()
                                    index_file_entry(
                                        conn, zname, zfname, zext,
                                        zsize, zmod, zcontent,
                                        inside_archive=zname,
                                        archive_path=filepath
                                    )
                                    indexing_state["indexed_count"] += 1
                                except Exception as ze:
                                    indexing_state["errors"].append(f"{zname} (in {fname}): {str(ze)[:80]}")
                    except Exception as ae:
                        indexing_state["errors"].append(f"{fname}: {str(ae)[:80]}")

                # Commit every 50 files
                if i % 50 == 0:
                    conn.commit()

            except Exception as e:
                indexing_state["errors"].append(f"{fname}: {str(e)[:80]}")
                indexing_state["skipped_count"] += 1

        conn.commit()

        # Register directory
        conn.execute("""
            INSERT OR REPLACE INTO indexed_dirs (path, indexed_at, file_count)
            VALUES (?, ?, ?)
        """, (directory, datetime.now().isoformat(), indexing_state["indexed_count"]))
        conn.commit()
        conn.close()

    except Exception as e:
        indexing_state["errors"].append(f"Erro crítico: {str(e)}")
    finally:
        indexing_state["running"] = False
        indexing_state["done"] = True
        indexing_state["current_file"] = "Concluído"


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route('/api/index', methods=['POST'])
def start_indexing():
    global indexing_state
    data = request.json
    directory = data.get('directory', '').strip()
    if not directory or not os.path.isdir(directory):
        return jsonify({"error": "Diretório inválido ou não encontrado"}), 400
    if indexing_state["running"]:
        return jsonify({"error": "Indexação já em andamento"}), 400
    t = threading.Thread(target=do_indexing, args=(directory,), daemon=True)
    t.start()
    return jsonify({"status": "started"})


@app.route('/api/stop', methods=['POST'])
def stop_indexing():
    global indexing_state
    indexing_state["running"] = False
    return jsonify({"status": "stopping"})


@app.route('/api/progress')
def get_progress():
    s = indexing_state
    elapsed = time.time() - s["start_time"] if s["start_time"] else 0
    rate = s["indexed_count"] / elapsed if elapsed > 0 else 0
    remaining = (s["total"] - s["progress"]) / rate if rate > 0 else 0
    return jsonify({
        "running": s["running"],
        "done": s["done"],
        "progress": s["progress"],
        "total": s["total"],
        "current_file": s["current_file"],
        "indexed_count": s["indexed_count"],
        "skipped_count": s["skipped_count"],
        "errors": s["errors"][-20:],
        "elapsed": round(elapsed),
        "rate": round(rate, 1),
        "remaining": round(remaining),
        "pct": round(s["progress"] / s["total"] * 100) if s["total"] > 0 else 0,
    })


@app.route('/api/search')
def search():
    query = request.args.get('q', '').strip()
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    ext_filter = request.args.get('ext', '')
    archive_filter = request.args.get('archive', '')

    if not query:
        return jsonify({"results": [], "total": 0})

    offset = (page - 1) * per_page

    # Build FTS query - wrap in quotes if it's a phrase with spaces
    fts_query = query
    if ' ' in query and not (query.startswith('"') and query.endswith('"')):
        fts_query = f'"{query}"'

    try:
        conn = get_db()
        
        # Count total
        where_extra = ""
        params_count = [fts_query]
        if ext_filter:
            where_extra += " AND f.extension = ?"
            params_count.append(ext_filter)
        if archive_filter == 'yes':
            where_extra += " AND f.archive_path IS NOT NULL"
        elif archive_filter == 'no':
            where_extra += " AND f.archive_path IS NULL"

        count_sql = f"""
            SELECT COUNT(*) as cnt
            FROM file_content fc
            JOIN files f ON f.path = fc.path
            WHERE file_content MATCH ?
            {where_extra}
        """
        total = conn.execute(count_sql, params_count).fetchone()["cnt"]

        # Fetch results with snippet
        params = [fts_query] + params_count[1:] + [per_page, offset]
        results_sql = f"""
            SELECT 
                f.id, f.path, f.filename, f.extension, f.size,
                f.modified_at, f.word_count, f.archive_path, f.inside_archive,
                snippet(file_content, 1, '<mark>', '</mark>', '…', 30) as snippet
            FROM file_content fc
            JOIN files f ON f.path = fc.path
            WHERE file_content MATCH ?
            {where_extra}
            ORDER BY rank
            LIMIT ? OFFSET ?
        """
        rows = conn.execute(results_sql, params).fetchall()
        conn.close()

        results = []
        for row in rows:
            results.append({
                "id": row["id"],
                "path": row["path"],
                "filename": row["filename"],
                "extension": row["extension"],
                "size": row["size"],
                "modified_at": row["modified_at"],
                "word_count": row["word_count"],
                "archive_path": row["archive_path"],
                "inside_archive": row["inside_archive"],
                "snippet": row["snippet"],
            })

        return jsonify({"results": results, "total": total, "page": page, "per_page": per_page})

    except Exception as e:
        return jsonify({"error": str(e), "results": [], "total": 0}), 500


@app.route('/api/stats')
def get_stats():
    try:
        conn = get_db()
        total_files = conn.execute("SELECT COUNT(*) as c FROM files").fetchone()["c"]
        total_indexed = conn.execute("SELECT COUNT(*) as c FROM file_content").fetchone()["c"]
        by_ext = conn.execute("""
            SELECT extension, COUNT(*) as cnt 
            FROM files 
            GROUP BY extension 
            ORDER BY cnt DESC 
            LIMIT 20
        """).fetchall()
        dirs = conn.execute("""
            SELECT path, indexed_at, file_count FROM indexed_dirs ORDER BY indexed_at DESC
        """).fetchall()
        total_size = conn.execute("SELECT SUM(size) as s FROM files").fetchone()["s"] or 0
        conn.close()
        return jsonify({
            "total_files": total_files,
            "total_indexed": total_indexed,
            "by_ext": [{"ext": r["extension"], "count": r["cnt"]} for r in by_ext],
            "dirs": [{"path": r["path"], "indexed_at": r["indexed_at"], "count": r["file_count"]} for r in dirs],
            "total_size": total_size,
        })
    except Exception as e:
        return jsonify({"error": str(e)})


@app.route('/api/clear', methods=['POST'])
def clear_index():
    try:
        conn = get_db()
        conn.executescript("DELETE FROM files; DELETE FROM file_content; DELETE FROM indexed_dirs;")
        conn.commit()
        conn.close()
        return jsonify({"status": "cleared"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/browse')
def browse_dirs():
    """Browse filesystem directories for the picker."""
    path = request.args.get('path', '').strip()

    # Default to home directory
    if not path:
        path = os.path.expanduser('~')

    path = os.path.expanduser(path)
    path = os.path.realpath(path)

    if not os.path.isdir(path):
        return jsonify({"error": "Caminho inválido"}), 400

    try:
        entries = []
        with os.scandir(path) as it:
            for entry in sorted(it, key=lambda e: (not e.is_dir(), e.name.lower())):
                if entry.name.startswith('.'):
                    continue
                if entry.is_dir(follow_symlinks=False):
                    entries.append({
                        "name": entry.name,
                        "path": entry.path,
                        "is_dir": True,
                    })

        # Build breadcrumbs
        parts = []
        p = path
        while True:
            parent = os.path.dirname(p)
            parts.append({"name": os.path.basename(p) or p, "path": p})
            if parent == p:
                break
            p = parent
        parts.reverse()

        return jsonify({
            "path": path,
            "parent": os.path.dirname(path) if path != os.path.dirname(path) else None,
            "entries": entries,
            "breadcrumbs": parts,
        })
    except PermissionError:
        return jsonify({"error": "Sem permissão para acessar este diretório"}), 403
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _reveal_in_file_manager(target: str) -> None:
    """Reveal a file/folder in the OS native file manager (cross-platform)."""
    system = platform.system()
    if system == 'Darwin':  # macOS
        subprocess.Popen(['open', '-R', target])
    elif system == 'Windows':
        # /select, highlights the file inside Explorer
        # Note: no space after the comma is required by explorer.exe
        subprocess.Popen(['explorer', f'/select,{target}'])
    else:  # Linux / *BSD — open the containing folder
        folder = target if os.path.isdir(target) else os.path.dirname(target)
        subprocess.Popen(['xdg-open', folder])


@app.route('/api/open', methods=['POST'])
def open_file():
    """Reveal file in the OS file manager (Finder / Explorer / xdg-open)."""
    data = request.json
    path = data.get('path', '')
    archive = data.get('archive_path', '')
    try:
        if archive and os.path.exists(archive):
            _reveal_in_file_manager(archive)
        elif os.path.exists(path):
            _reveal_in_file_manager(path)
        return jsonify({"status": "ok"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─── HTML Template ─────────────────────────────────────────────────────────────
_UI_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ui.html')
with open(_UI_PATH, 'r', encoding='utf-8') as _f:
    HTML_TEMPLATE = _f.read()


if __name__ == '__main__':
    init_db()
    print("\n" + "="*50)
    print("  📂 File Indexer — Iniciando...")
    print("  Abra: http://localhost:7432")
    print("="*50 + "\n")
    import webbrowser
    threading.Timer(1.0, lambda: webbrowser.open("http://localhost:7432")).start()
    app.run(host='127.0.0.1', port=7432, debug=False, threaded=True)
